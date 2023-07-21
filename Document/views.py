import os
from django.shortcuts import render
from django.db.models import Q
from django.conf import settings
from django.utils.dateparse import parse_date
from django.http import HttpResponse, Http404
from django.shortcuts import get_object_or_404
from django.contrib.auth import get_user_model
User = get_user_model()

from Account.models import Documents
from Document.serializers import DocumentsSerializers, DocumentsSharingSerializers

from rest_framework import generics
from rest_framework.permissions import IsAuthenticated, IsAdminUser
from rest_framework.authentication import BasicAuthentication
from rest_framework_simplejwt.authentication import JWTAuthentication
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.views import APIView
from rest_framework import status
from rest_framework.response import Response
from rest_framework.filters import SearchFilter
from rest_framework.exceptions import PermissionDenied


# from PyPDF2 import PdfWriter, PdfReader
import PyPDF4

from docx2pdf import convert
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import pdfplumber



# Create your views here.
##-----------------------------( User )------------------------------------------------------

# URL = ( http://127.0.0.1:8000/document/user/document-uplode/ )
class UserDocumentsUplodeAPIView(APIView):
    parser_classes = [MultiPartParser, FormParser]
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        file = request.FILES.get('file', None)

        if not file:
            return Response({"error": "No file was submitted."}, status=status.HTTP_400_BAD_REQUEST)

        file_format = request.data.get('file_format', '').lower()

        if file_format not in ['pdf', 'docx', 'txt']:
            return Response({"error": "Invalid file format. Only 'pdf', 'docx', and 'txt' are supported."},
                            status=status.HTTP_400_BAD_REQUEST)

        try:
            # Save the uploaded file to the Documents model
            document = Documents.objects.create(user=request.user, title=request.data.get('title', ''),
                                                description=request.data.get('description', ''), file=file,
                                                file_format=file_format)

            # Convert the uploaded file to the desired format
            if file_format == 'pdf':
                if file.name.lower().endswith('.pdf'):
                    # If the uploaded file is already a PDF, save it directly
                    document.save()
                else:
                    # Convert to PDF
                    pdf_buffer = BytesIO()
                    c = canvas.Canvas(pdf_buffer, pagesize=letter)
                    if file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                        # Convert DOCX to PDF
                        doc = Document(document.file)
                        for para in doc.paragraphs:
                            c.drawString(72, 800, para.text)
                            c.showPage()
                    elif file.content_type == 'text/plain':
                        # Convert TXT to PDF
                        text = document.file.read().decode('utf-8')
                        c.drawString(72, 800, text)
                        c.showPage()
                    # Add more conditions for other file types if needed.
                    c.save()
                    pdf_buffer.seek(0)
                    new_file = f'{os.path.splitext(file.name)[0]}.pdf'
                    path = os.path.join(settings.MEDIA_ROOT, new_file)
                    with open(path, 'wb') as f:
                        f.write(pdf_buffer.read())
                    document.file = new_file
                    document.save()

            elif file_format == 'docx':
                if file.name.lower().endswith('.docx'):
                    # If the uploaded file is already a DOCX, save it directly
                    document.save()
                else:
                    # Convert to DOCX
                    doc = Document()
                    if file.content_type == 'application/pdf':
                        # Convert PDF to DOCX
                        pdf = pdfplumber.open(document.file)
                        text = ''
                        for page in pdf.pages:
                            text += page.extract_text()
                        pdf.close()
                        doc.add_paragraph(text)
                    elif file.content_type == 'text/plain':
                        # Convert TXT to DOCX
                        text = document.file.read().decode('utf-8')
                        doc.add_paragraph(text)
                    # Add more conditions for other file types if needed.
                    new_file = f'{os.path.splitext(file.name)[0]}.docx'
                    path = os.path.join(settings.MEDIA_ROOT, new_file)
                    doc.save(path)
                    document.file = new_file
                    document.save()

            elif file_format == 'txt':
                if file.name.lower().endswith('.txt'):
                    # If the uploaded file is already a TXT, save it directly
                    document.save()
                # Convert to TXT (if the file is not already a TXT)
                elif file.content_type != 'text/plain':
                    new_file = f'{os.path.splitext(file.name)[0]}.txt'
                    path = os.path.join(settings.MEDIA_ROOT, new_file)
                    with open(path, 'w', encoding='utf-8', errors='ignore') as f:
                        f.write(document.file.read().decode('utf-8', errors='ignore'))
                    document.file = new_file
                    document.save()

            # Serialize and return the converted file details
            serializer = DocumentsSerializers(document)
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



# URL = ( http://127.0.0.1:8000/document/user/document-display/)
class UserDocumentsListAPIView(generics.ListAPIView):
    authentication_classes = [JWTAuthentication]
    # authentication_classes = [BasicAuthentication]
    permission_classes = [IsAuthenticated]
    serializer_class = DocumentsSharingSerializers

    filter_backends = [ SearchFilter ]
    search_fields = {
        'user__first_name': ['icontains'],
        'user__last_name': ['icontains'],
        'title': ['icontains'],
        'description': ['icontains'],
        'file_format': ['exact'],
        # 'create_at': ['exact'],
        'create_at': ['date'],
    }

    def get_queryset(self):
        user = self.request.user  # Get the current user
        if self.request.user.is_staff == False:
            return Documents.objects.filter(user=user)
        else:
            return Documents.objects.all()


# URL ( http://127.0.0.1:8000/document/ user/documents-sharing/<str:title>/ )
class DocumentsSharingDetailView(generics.RetrieveAPIView):
    authentication_classes = [JWTAuthentication]
    # authentication_classes = [BasicAuthentication]
    permission_classes = [IsAuthenticated]
    queryset = Documents.objects.all()
    serializer_class = DocumentsSharingSerializers
    lookup_field = 'title'





# URL = ( http://127.0.0.1:8000/document/user/document-retrieve-update-delete/<int:pk>/)
class UserDocumentsRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateDestroyAPIView):
    authentication_classes = [JWTAuthentication]
    # authentication_classes = [BasicAuthentication]
    permission_classes = [IsAuthenticated]
    serializer_class = DocumentsSerializers

    def get_queryset(self):
        user = self.request.user  # Get the current user
        if user.is_staff == True:
            doc_id = self.kwargs['pk']
            return Documents.objects.filter(pk = doc_id)
        else:
            return Documents.objects.filter(user=user)
    





##-----------------------------( Admin Pannel )------------------------------------------------------

# URL = ( http://127.0.0.1:8000/document/admin/document-retrieve/ )
class AdminDocumentsListAPIView(generics.ListAPIView):
    queryset = Documents.objects.all()
    # authentication_classes = [JWTAuthentication]
    authentication_classes = [BasicAuthentication]
    permission_classes = [IsAuthenticated, IsAdminUser]
    serializer_class = DocumentsSharingSerializers

    filter_backends = [ SearchFilter ]
    search_fields = {
        'user__first_name': ['icontains'],
        'user__last_name': ['icontains'],
        'title': ['icontains'],
        'description': ['icontains'],
        'file_format': ['exact'],
        # 'create_at': ['exact'],
        'create_at': ['date'],
    }


    def get_queryset(self):
        queryset = super().get_queryset()

        # Custom handling for create_at search
        create_at_search = self.request.query_params.get('create_at', None)
        if create_at_search:
            # Parse the date string to a Python date object
            create_at_date = parse_date(create_at_search)
            if create_at_date:
                # Filter the queryset to get Documents with the specified create_at date
                queryset = queryset.filter(create_at__date=create_at_date)

        return queryset




# URL = ( http://127.0.0.1:8000/document/admin/document-retrieve-update-delete/<int:pk>/)
class AdminDocumentsRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateDestroyAPIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsAdminUser]
    serializer_class = DocumentsSerializers

    def get_queryset(self):
        user = self.request.user  # Get the current user
        if user.is_staff == True:
            doc_id = self.kwargs['pk']
            return Documents.objects.filter(pk = doc_id)
        







# URL = ( http://127.0.0.1:8000/document/download-document/<int:file_id>/<str:doc_format>/ )
class DownloadFileView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request, file_id, doc_format):
        document = get_object_or_404(Documents, id=file_id)

        # Check if the authenticated user is the owner of the file or a staff user
        if not (document.user == request.user or request.user.is_staff):
            raise PermissionDenied("You are not allowed to download this file.")

        # Validate that the requested format is either 'pdf', 'docx', or 'txt'
        if doc_format not in ['pdf', 'docx', 'txt']:
            return HttpResponse("Invalid format. Only 'pdf', 'docx', and 'txt' formats are supported.", status=400)

        # Get the file path and content type based on the requested format
        file_path = document.file.path
        content_type = f'application/{doc_format}'

        # Set the appropriate response headers for the file download
        response = HttpResponse(open(file_path, 'rb'), content_type=content_type)
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'

        return response
